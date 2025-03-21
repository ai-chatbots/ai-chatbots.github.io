from fastapi import APIRouter, UploadFile, File, HTTPException
import os
import fitz  # PyMuPDF for PDFs
import docx  # python-docx for DOCX files
from io import BytesIO
import zipfile

from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from app.config import OPENAI_API_KEY

router = APIRouter()

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    text = ""
    try:
        file_stream = BytesIO(file_content)
        pdf_document = fitz.open(stream=file_stream.read(), filetype="pdf")
        for page in pdf_document:
            text += page.get_text()
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {e}")
    return text

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        file_stream = BytesIO(file_content)
        # Validate that the file is a proper DOCX (which is a zip file)
        if not zipfile.is_zipfile(file_stream):
            raise Exception("File is not a valid DOCX file (zip file expected).")
        file_stream.seek(0)
        document = docx.Document(file_stream)
        text = "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        raise Exception(f"Error extracting DOCX text: {e}")
    return text

@router.post("/upload-documents")
async def upload_documents(file: UploadFile = File(...)):
    # Supported file types
    supported_types = [
        "text/plain",
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    if file.content_type not in supported_types:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Supported types: .txt, .pdf, .docx"
        )

    content = await file.read()
    # Determine text extraction based on content type.
    if file.content_type == "text/plain":
        text = content.decode("utf-8")
    elif file.content_type == "application/pdf":
        try:
            text = extract_text_from_pdf(content)
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            text = extract_text_from_docx(content)
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    # Create a Document object from the extracted text.
    doc = Document(page_content=text)

    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    docs = text_splitter.split_documents([doc])

    # Initialize embeddings using OpenAI API key from your config.
    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)

    # Build a FAISS vector store from the uploaded documents.
    vector_store = FAISS.from_documents(docs, embeddings)

    # Save the FAISS index using its built-in save method.
    try:
        vector_store.save_local("faiss_index")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving FAISS index: {str(e)}")

    return {"detail": "Documents uploaded successfully and FAISS index updated."}
